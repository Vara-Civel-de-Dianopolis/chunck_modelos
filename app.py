import os
import re
import mysql.connector
from mysql.connector import Error
from typing import List, Dict, Tuple, Optional
import PyPDF2
import docx
from datetime import datetime
import logging
import hashlib
from pathlib import Path
import time
import json
import google.generativeai as genai

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class LegalDocumentChapterDetector:
    """Classe para detectar e separar capítulos em documentos jurídicos usando Gemini AI"""
    
    def __init__(self, api_key: str):
        """
        Inicializa o detector de capítulos
        
        Args:
            api_key: Chave da API do Google Gemini
        """
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')
        
    def detect_chapters(self, text: str, document_name: str) -> List[Dict]:
        """
        Detecta capítulos em documentos jurídicos usando IA
        
        Args:
            text: Texto completo do documento
            document_name: Nome do documento para contexto
            
        Returns:
            Lista de capítulos com suas posições e conteúdo
        """
        try:
            # Prompt especializado para documentos jurídicos
            prompt = f"""
            Analise o seguinte documento jurídico e identifique TODOS os capítulos, seções e subdivisões estruturais.
            
            DOCUMENTO: {document_name}
            
            INSTRUÇÕES:
            1. Identifique a estrutura hierárquica completa (Capítulos, Seções, Subseções, etc.)
            2. Para cada divisão encontrada, forneça:
               - Título completo
               - Posição aproximada no texto (caractere inicial)
               - Tipo (CAPITULO, SECAO, SUBSECAO, etc.)
               - Nível hierárquico (1, 2, 3, etc.)
            
            3. Considere padrões comuns em documentos jurídicos:
               - "CAPÍTULO I", "CAPÍTULO II", etc.
               - "Seção I", "Seção II", etc.
               - "Art.", "Artigo"
               - "1.", "2.", "3." (numeração)
               - "a)", "b)", "c)" (alíneas)
               - "I -", "II -", "III -" (numeração romana)
               - Títulos em maiúsculas
               - "DISPOSIÇÕES GERAIS", "DISPOSIÇÕES FINAIS"
               - "CONSIDERANDO", "FUNDAMENTAÇÃO"
               - "DISPOSITIVO", "CONCLUSÃO"
            
            4. Retorne APENAS um JSON válido no formato:
            {{
                "chapters": [
                    {{
                        "title": "título completo",
                        "type": "CAPITULO|SECAO|SUBSECAO|ARTIGO|DISPOSITIVO",
                        "level": 1,
                        "start_position": 0,
                        "content_preview": "primeiras 100 palavras do conteúdo"
                    }}
                ]
            }}
            
            TEXTO DO DOCUMENTO:
            {text[:8000]}...
            
            IMPORTANTE: Retorne APENAS o JSON, sem explicações adicionais.
            """
            
            response = self.model.generate_content(prompt)
            
            # Tenta extrair JSON da resposta
            response_text = response.text.strip()
            
            # Remove possíveis marcadores de código
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            
            # Parse do JSON
            try:
                result = json.loads(response_text)
                chapters = result.get('chapters', [])
                
                # Valida e ajusta as posições dos capítulos
                validated_chapters = self._validate_and_adjust_chapters(chapters, text)
                
                logger.info(f"Detectados {len(validated_chapters)} capítulos no documento {document_name}")
                return validated_chapters
                
            except json.JSONDecodeError as e:
                logger.warning(f"Erro ao parsear JSON da resposta do Gemini: {e}")
                logger.debug(f"Resposta recebida: {response_text[:500]}...")
                
                # Fallback: detecção baseada em regras
                return self._fallback_chapter_detection(text)
                
        except Exception as e:
            logger.error(f"Erro na detecção de capítulos com Gemini: {e}")
            # Fallback: detecção baseada em regras
            return self._fallback_chapter_detection(text)
    
    def _validate_and_adjust_chapters(self, chapters: List[Dict], text: str) -> List[Dict]:
        """
        Valida e ajusta as posições dos capítulos detectados
        
        Args:
            chapters: Lista de capítulos detectados
            text: Texto completo do documento
            
        Returns:
            Lista de capítulos validados
        """
        validated_chapters = []
        text_lower = text.lower()
        
        for chapter in chapters:
            try:
                title = chapter.get('title', '').strip()
                if not title:
                    continue
                
                # Busca a posição real do título no texto
                title_lower = title.lower()
                
                # Tenta diferentes variações do título
                search_variations = [
                    title_lower,
                    title_lower.replace(' ', ''),
                    re.sub(r'[^\w\s]', '', title_lower),
                    title_lower.split()[0] if title_lower.split() else ''
                ]
                
                found_position = None
                for variation in search_variations:
                    if variation and len(variation) > 3:
                        pos = text_lower.find(variation)
                        if pos != -1:
                            found_position = pos
                            break
                
                if found_position is not None:
                    chapter['start_position'] = found_position
                    chapter['actual_title'] = title
                    
                    # Extrai preview do conteúdo
                    preview_end = min(found_position + 500, len(text))
                    chapter['content_preview'] = text[found_position:preview_end].strip()
                    
                    validated_chapters.append(chapter)
                else:
                    logger.debug(f"Não foi possível localizar o capítulo: {title}")
                    
            except Exception as e:
                logger.warning(f"Erro ao validar capítulo {chapter}: {e}")
                continue
        
        # Ordena capítulos por posição
        validated_chapters.sort(key=lambda x: x['start_position'])
        
        return validated_chapters
    
    def _fallback_chapter_detection(self, text: str) -> List[Dict]:
        """
        Detecção de capítulos baseada em regras como fallback
        
        Args:
            text: Texto do documento
            
        Returns:
            Lista de capítulos detectados
        """
        logger.info("Usando detecção de capítulos baseada em regras (fallback)")
        
        chapters = []
        
        # Padrões comuns em documentos jurídicos
        patterns = [
            # Capítulos
            (r'(?i)^(CAPÍTULO\s+[IVX]+|CAPÍTULO\s+\d+)[\s\-–—]*(.{0,100}?)(?=\n|\r|$)', 'CAPITULO', 1),
            (r'(?i)^(CAP\.?\s+[IVX]+|CAP\.?\s+\d+)[\s\-–—]*(.{0,100}?)(?=\n|\r|$)', 'CAPITULO', 1),
            
            # Seções
            (r'(?i)^(SEÇÃO\s+[IVX]+|SEÇÃO\s+\d+)[\s\-–—]*(.{0,100}?)(?=\n|\r|$)', 'SECAO', 2),
            (r'(?i)^(SEÇ\.?\s+[IVX]+|SEÇ\.?\s+\d+)[\s\-–—]*(.{0,100}?)(?=\n|\r|$)', 'SECAO', 2),
            
            # Artigos
            (r'(?i)^(ART\.?\s+\d+|ARTIGO\s+\d+)[\s\-–—]*(.{0,100}?)(?=\n|\r|$)', 'ARTIGO', 3),
            
            # Disposições especiais
            (r'(?i)^(DISPOSIÇÕES?\s+GERAIS|DISPOSIÇÕES?\s+FINAIS|DISPOSIÇÕES?\s+TRANSITÓRIAS)(.{0,50}?)(?=\n|\r|$)', 'DISPOSITIVO', 1),
            
            # Considerandos e fundamentação
            (r'(?i)^(CONSIDERANDO|FUNDAMENTAÇÃO|MOTIVAÇÃO)(.{0,50}?)(?=\n|\r|$)', 'FUNDAMENTACAO', 1),
            
            # Dispositivo e conclusão
            (r'(?i)^(DISPOSITIVO|CONCLUSÃO|DECISÃO)(.{0,50}?)(?=\n|\r|$)', 'DISPOSITIVO', 1),
            
            # Títulos em maiúsculas (genérico)
            (r'^([A-ZÁÉÍÓÚÂÊÎÔÛÀÈÌÒÙÃÕÇ\s]{10,80})(?=\n|\r|$)', 'TITULO', 2),
        ]
        
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            for pattern, chapter_type, level in patterns:
                match = re.match(pattern, line, re.MULTILINE)
                if match:
                    # Calcula posição no texto original
                    start_pos = sum(len(lines[j]) + 1 for j in range(i))
                    
                    title = match.group(1)
                    if len(match.groups()) > 1 and match.group(2):
                        title += " " + match.group(2).strip()
                    
                    title = title.strip()
                    
                    # Evita duplicatas próximas
                    if not any(abs(ch['start_position'] - start_pos) < 50 for ch in chapters):
                        chapter = {
                            'title': title,
                            'type': chapter_type,
                            'level': level,
                            'start_position': start_pos,
                            'content_preview': line[:200]
                        }
                        chapters.append(chapter)
                    break
        
        # Ordena por posição
        chapters.sort(key=lambda x: x['start_position'])
        
        logger.info(f"Detecção por regras encontrou {len(chapters)} capítulos")
        return chapters

class LegalDocumentChunker:
    """Classe especializada para chunking de documentos jurídicos por capítulos"""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        """
        Inicializa o chunker para documentos jurídicos
        
        Args:
            chunk_size: Tamanho de cada chunk em caracteres
            overlap: Sobreposição entre chunks em caracteres
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
        
    def create_chapter_chunks(self, text: str, chapters: List[Dict]) -> List[Dict]:
        """
        Cria chunks organizados por capítulos
        
        Args:
            text: Texto completo do documento
            chapters: Lista de capítulos detectados
            
        Returns:
            Lista de chunks organizados por capítulos
        """
        all_chunks = []
        
        if not chapters:
            # Se não há capítulos, trata como um único capítulo
            logger.warning("Nenhum capítulo detectado, processando como documento único")
            chapter_chunks = self._create_chunks_for_chapter(
                text, 
                {
                    'title': 'DOCUMENTO COMPLETO',
                    'type': 'DOCUMENTO',
                    'level': 1,
                    'start_position': 0
                },
                0,
                len(text)
            )
            all_chunks.extend(chapter_chunks)
            return all_chunks
        
        # Processa cada capítulo
        for i, chapter in enumerate(chapters):
            start_pos = chapter['start_position']
            
            # Determina fim do capítulo (início do próximo ou fim do documento)
            if i + 1 < len(chapters):
                end_pos = chapters[i + 1]['start_position']
            else:
                end_pos = len(text)
            
            # Extrai texto do capítulo
            chapter_text = text[start_pos:end_pos].strip()
            
            if chapter_text:
                chapter_chunks = self._create_chunks_for_chapter(
                    chapter_text, 
                    chapter, 
                    start_pos,
                    end_pos
                )
                all_chunks.extend(chapter_chunks)
        
        return all_chunks
    
    def _create_chunks_for_chapter(self, chapter_text: str, chapter_info: Dict, 
                                  chapter_start: int, chapter_end: int) -> List[Dict]:
        """
        Cria chunks para um capítulo específico
        
        Args:
            chapter_text: Texto do capítulo
            chapter_info: Informações do capítulo
            chapter_start: Posição inicial do capítulo no documento
            chapter_end: Posição final do capítulo no documento
            
        Returns:
            Lista de chunks do capítulo
        """
        chunks = []
        
        if len(chapter_text) <= self.chunk_size:
            # Capítulo cabe em um único chunk
            chunk = {
                'chunk_index': 0,
                'content': chapter_text,
                'chunk_size': len(chapter_text),
                'start_position': 0,
                'end_position': len(chapter_text),
                'overlap_size': 0,
                'chapter_title': chapter_info['title'],
                'chapter_type': chapter_info['type'],
                'chapter_level': chapter_info['level'],
                'absolute_start_position': chapter_start,
                'absolute_end_position': chapter_end,
                'is_chapter_complete': True
            }
            chunks.append(chunk)
            return chunks
        
        # Divide capítulo em múltiplos chunks
        start = 0
        chunk_index = 0
        
        while start < len(chapter_text):
            end = min(start + self.chunk_size, len(chapter_text))
            
            # Tenta quebrar em uma frase completa para documentos jurídicos
            chunk_content = chapter_text[start:end]
            
            if end < len(chapter_text):
                # Procura por quebras naturais em documentos jurídicos
                break_points = [
                    chunk_content.rfind('.'),
                    chunk_content.rfind(';'),
                    chunk_content.rfind(':'),
                    chunk_content.rfind('\n'),
                    chunk_content.rfind(' ')
                ]
                
                best_break = max([bp for bp in break_points if bp > self.chunk_size * 0.7])
                if best_break > 0:
                    end = start + best_break + 1
                    chunk_content = chapter_text[start:end]
            
            chunk = {
                'chunk_index': chunk_index,
                'content': chunk_content.strip(),
                'chunk_size': len(chunk_content.strip()),
                'start_position': start,
                'end_position': end,
                'overlap_size': self.overlap if chunk_index > 0 else 0,
                'chapter_title': chapter_info['title'],
                'chapter_type': chapter_info['type'],
                'chapter_level': chapter_info['level'],
                'absolute_start_position': chapter_start + start,
                'absolute_end_position': chapter_start + end,
                'is_chapter_complete': False
            }
            
            chunks.append(chunk)
            
            # Move para próxima posição com overlap
            start = end - self.overlap
            chunk_index += 1
            
            # Evita chunks muito pequenos no final
            if len(chapter_text) - start < self.chunk_size * 0.3:
                break
        
        return chunks

class SmartLocalFileProcessor:
    """Classe inteligente para processar documentos locais com detecção de mudanças"""
    
    def __init__(self, base_folder: str):
        """
        Inicializa o processador de arquivos locais
        
        Args:
            base_folder: Caminho para a pasta base dos documentos
        """
        self.base_folder = Path(base_folder)
        if not self.base_folder.exists():
            raise ValueError(f"Pasta não encontrada: {base_folder}")
        
        # Estatísticas de busca
        self.search_stats = {
            'folders_scanned': 0,
            'files_found': 0,
            'files_by_type': {},
            'folders_with_errors': []
        }
            
    def list_files(self, file_types: List[str] = None, 
                   recursive: bool = True, 
                   max_depth: Optional[int] = None,
                   show_progress: bool = True) -> List[Dict]:
        """
        Lista arquivos da pasta local com busca recursiva aprimorada
        
        Args:
            file_types: Lista de extensões de arquivo (ex: ['pdf', 'docx', 'txt'])
            recursive: Se deve buscar em subpastas
            max_depth: Profundidade máxima de busca (None = sem limite)
            show_progress: Se deve mostrar progresso da busca
            
        Returns:
            Lista de informações dos arquivos
        """
        if file_types is None:
            file_types = ['pdf', 'docx', 'txt', 'rtf', 'odt']
            
        # Normaliza as extensões (remove pontos e converte para minúsculo)
        file_types = [ext.lower().lstrip('.') for ext in file_types]
        
        # Reset das estatísticas
        self.search_stats = {
            'folders_scanned': 0,
            'files_found': 0,
            'files_by_type': {ext: 0 for ext in file_types},
            'folders_with_errors': []
        }
        
        files = []
        
        if recursive:
            files = self._scan_recursive(
                self.base_folder, 
                file_types, 
                max_depth, 
                show_progress
            )
        else:
            files = self._scan_single_folder(self.base_folder, file_types)
            
        # Log das estatísticas finais
        self._log_search_stats()
        
        return files
    
    def _scan_recursive(self, folder: Path, file_types: List[str], 
                       max_depth: Optional[int], show_progress: bool, 
                       current_depth: int = 0) -> List[Dict]:
        """
        Escaneia recursivamente todas as subpastas
        
        Args:
            folder: Pasta atual
            file_types: Tipos de arquivo aceitos
            max_depth: Profundidade máxima
            show_progress: Mostrar progresso
            current_depth: Profundidade atual
            
        Returns:
            Lista de arquivos encontrados
        """
        files = []
        
        try:
            # Verifica limite de profundidade
            if max_depth is not None and current_depth > max_depth:
                return files
            
            self.search_stats['folders_scanned'] += 1
            
            if show_progress and self.search_stats['folders_scanned'] % 10 == 0:
                logger.info(f"Escaneadas {self.search_stats['folders_scanned']} pastas, "
                          f"encontrados {self.search_stats['files_found']} arquivos...")
            
            # Lista todos os itens da pasta atual
            try:
                items = list(folder.iterdir())
            except PermissionError:
                logger.warning(f"Sem permissão para acessar: {folder}")
                self.search_stats['folders_with_errors'].append(str(folder))
                return files
            except Exception as e:
                logger.error(f"Erro ao acessar pasta {folder}: {e}")
                self.search_stats['folders_with_errors'].append(str(folder))
                return files
            
            # Separa arquivos e pastas
            current_files = []
            subfolders = []
            
            for item in items:
                if item.is_file():
                    current_files.append(item)
                elif item.is_dir():
                    subfolders.append(item)
            
            # Processa arquivos da pasta atual
            folder_files = self._process_files_in_folder(current_files, file_types, folder)
            files.extend(folder_files)
            
            # Processa subpastas recursivamente
            for subfolder in subfolders:
                try:
                    subfolder_files = self._scan_recursive(
                        subfolder, 
                        file_types, 
                        max_depth, 
                        show_progress, 
                        current_depth + 1
                    )
                    files.extend(subfolder_files)
                except Exception as e:
                    logger.error(f"Erro ao processar subpasta {subfolder}: {e}")
                    self.search_stats['folders_with_errors'].append(str(subfolder))
                    
        except Exception as e:
            logger.error(f"Erro geral ao escanear pasta {folder}: {e}")
            self.search_stats['folders_with_errors'].append(str(folder))
            
        return files
    
    def _scan_single_folder(self, folder: Path, file_types: List[str]) -> List[Dict]:
        """
        Escaneia apenas uma pasta (sem recursão)
        
        Args:
            folder: Pasta para escanear
            file_types: Tipos de arquivo aceitos
            
        Returns:
            Lista de arquivos encontrados
        """
        files = []
        self.search_stats['folders_scanned'] = 1
        
        try:
            items = [item for item in folder.iterdir() if item.is_file()]
            files = self._process_files_in_folder(items, file_types, folder)
        except Exception as e:
            logger.error(f"Erro ao escanear pasta {folder}: {e}")
            self.search_stats['folders_with_errors'].append(str(folder))
            
        return files
    
    def _process_files_in_folder(self, files: List[Path], 
                                file_types: List[str], 
                                folder: Path) -> List[Dict]:
        """
        Processa arquivos de uma pasta específica
        
        Args:
            files: Lista de arquivos da pasta
            file_types: Tipos de arquivo aceitos
            folder: Pasta atual
            
        Returns:
            Lista de informações dos arquivos válidos
        """
        valid_files = []
        
        for file_path in files:
            try:
                file_extension = file_path.suffix.lower().lstrip('.')
                
                if file_extension in file_types:
                    # Obtém informações do arquivo
                    file_stat = file_path.stat()
                    
                    # Calcula caminho relativo
                    try:
                        relative_path = file_path.relative_to(self.base_folder)
                    except ValueError:
                        # Se não conseguir calcular relativo, usa o absoluto
                        relative_path = file_path
                    
                    file_info = {
                        'name': file_path.name,
                        'path': str(file_path.absolute()),
                        'relative_path': str(relative_path),
                        'folder': str(folder),
                        'size': file_stat.st_size,
                        'extension': file_extension,
                        'modified_time': datetime.fromtimestamp(file_stat.st_mtime),
                        'created_time': datetime.fromtimestamp(file_stat.st_ctime),
                        'depth': len(relative_path.parts) - 1,  # Profundidade da pasta
                        # Informações para detecção de mudanças
                        'modification_timestamp': file_stat.st_mtime,
                        'file_hash': self._calculate_file_hash(file_path)
                    }
                    
                    valid_files.append(file_info)
                    self.search_stats['files_found'] += 1
                    self.search_stats['files_by_type'][file_extension] += 1
                    
            except Exception as e:
                logger.warning(f"Erro ao processar arquivo {file_path}: {e}")
                continue
                
        return valid_files
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """
        Calcula hash MD5 do arquivo para detecção de mudanças
        
        Args:
            file_path: Caminho do arquivo
            
        Returns:
            Hash MD5 do arquivo
        """
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                # Lê o arquivo em chunks para não sobrecarregar a memória
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.warning(f"Erro ao calcular hash do arquivo {file_path}: {e}")
            # Fallback: usa informações do arquivo
            file_stat = file_path.stat()
            fallback_string = f"{file_path.absolute()}_{file_stat.st_size}_{file_stat.st_mtime}"
            return hashlib.md5(fallback_string.encode()).hexdigest()
    
    def _log_search_stats(self):
        """Log das estatísticas de busca"""
        logger.info("="*60)
        logger.info("ESTATÍSTICAS DA BUSCA DE ARQUIVOS")
        logger.info("="*60)
        logger.info(f"Pastas escaneadas: {self.search_stats['folders_scanned']}")
        logger.info(f"Total de arquivos encontrados: {self.search_stats['files_found']}")
        
        if self.search_stats['files_by_type']:
            logger.info("Arquivos por tipo:")
            for file_type, count in self.search_stats['files_by_type'].items():
                if count > 0:
                    logger.info(f"  - .{file_type}: {count} arquivos")
        
        if self.search_stats['folders_with_errors']:
            logger.warning(f"Pastas com erro de acesso: {len(self.search_stats['folders_with_errors'])}")
            for folder in self.search_stats['folders_with_errors'][:5]:  # Mostra apenas as primeiras 5
                logger.warning(f"  - {folder}")
            if len(self.search_stats['folders_with_errors']) > 5:
                logger.warning(f"  ... e mais {len(self.search_stats['folders_with_errors']) - 5} pastas")
        
        logger.info("="*60)

class EnhancedTextExtractor:
    """Classe aprimorada para extrair texto de diferentes tipos de arquivo"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extrai texto de arquivo PDF com tratamento robusto de erros"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text = ""
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Erro ao extrair texto da página {page_num + 1} do PDF {file_path}: {e}")
                        continue
                
                if not text.strip():
                    logger.warning(f"Nenhum texto extraído do PDF {file_path}")
                    
                return text.strip()
                
        except Exception as e:
            logger.error(f"Erro ao extrair texto do PDF {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extrai texto de arquivo DOCX incluindo tabelas e cabeçalhos"""
        try:
            doc = docx.Document(file_path)
            
            text = ""
            
            # Extrai texto dos parágrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extrai texto das tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            # Extrai texto de cabeçalhos e rodapés
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                            
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                    
            return text.strip()
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_from_txt(file_path: str) -> str:
        """Extrai texto de arquivo TXT com detecção automática de encoding"""
        try:
            # Lista de encodings para tentar
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1', 'ascii']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content:  # Se conseguiu ler e tem conteúdo
                            logger.debug(f"Arquivo {file_path} lido com encoding {encoding}")
                            return content
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.warning(f"Erro ao ler {file_path} com encoding {encoding}: {e}")
                    continue
                    
            logger.error(f"Não foi possível decodificar o arquivo {file_path} com nenhum encoding testado")
            return ""
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do TXT {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_text(file_path: str, file_extension: str) -> str:
        """
        Método unificado para extrair texto baseado na extensão
        
        Args:
            file_path: Caminho do arquivo
            file_extension: Extensão do arquivo
            
        Returns:
            Texto extraído
        """
        file_extension = file_extension.lower()
        
        if file_extension == 'pdf':
            return EnhancedTextExtractor.extract_from_pdf(file_path)
        elif file_extension in ['docx', 'doc']:
            return EnhancedTextExtractor.extract_from_docx(file_path)
        elif file_extension == 'txt':
            return EnhancedTextExtractor.extract_from_txt(file_path)
        else:
            logger.warning(f"Tipo de arquivo não suportado: {file_extension}")
            return ""

class LegalDatabaseManager:
    """Classe especializada para gerenciar banco de dados de documentos jurídicos"""
    
    def __init__(self, host: str, database: str, user: str, password: str, port: int = 3306):
        """
        Inicializa o gerenciador do banco de dados
        
        Args:
            host: Host do MySQL
            database: Nome do banco de dados
            user: Usuário do MySQL
            password: Senha do MySQL
            port: Porta do MySQL
        """
        self.host = host
        self.database = database
        self.user = user
        self.password = password
        self.port = port
        self.connection = None
        
    def connect(self):
        """Conecta ao banco de dados MySQL"""
        try:
            self.connection = mysql.connector.connect(
                host=self.host,
                database=self.database,
                user=self.user,
                password=self.password,
                port=self.port,
                autocommit=False
            )
            logger.info("Conexão com MySQL estabelecida")
            
        except Error as e:
            logger.error(f"Erro ao conectar ao MySQL: {e}")
            raise
            
    def disconnect(self):
        """Desconecta do banco de dados"""
        if self.connection and self.connection.is_connected():
            self.connection.close()
            logger.info("Conexão com MySQL encerrada")
    
    def check_file_status(self, file_info: Dict) -> Dict:
        """
        Verifica o status do arquivo no banco de dados
        
        Args:
            file_info: Informações do arquivo
            
        Returns:
            Dicionário com status do arquivo
        """
        try:
            cursor = self.connection.cursor()
            
            # Busca arquivo pelo caminho
            query = """
            SELECT id, file_hash, modification_timestamp, last_processed, content_length
            FROM documents 
            WHERE file_path = %s
            """
            
            cursor.execute(query, (file_info['path'],))
            result = cursor.fetchone()
            
            if result:
                document_id, stored_hash, stored_timestamp, last_processed, content_length = result
                
                # Verifica se houve mudanças
                current_hash = file_info['file_hash']
                current_timestamp = file_info['modification_timestamp']
                
                needs_update = (
                    stored_hash != current_hash or 
                    abs(stored_timestamp - current_timestamp) > 1
                )
                
                return {
                    'exists': True,
                    'needs_update': needs_update,
                    'document_id': document_id,
                    'last_hash': stored_hash,
                    'last_modified': datetime.fromtimestamp(stored_timestamp),
                    'content_length': content_length
                }
            else:
                return {
                    'exists': False,
                    'needs_update': True,
                    'document_id': None,
                    'last_hash': None,
                    'last_modified': None,
                    'content_length': 0
                }
                
        except Error as e:
            logger.error(f"Erro ao verificar status do arquivo: {e}")
            return {
                'exists': False,
                'needs_update': True,
                'document_id': None,
                'last_hash': None,
                'last_modified': None,
                'content_length': 0
            }
        finally:
            if cursor:
                cursor.close()
            
    def insert_or_update_document(self, file_info: Dict, chapters: List[Dict], 
                                 document_id: Optional[int] = None) -> Optional[int]:
        """
        Insere ou atualiza informações do documento na tabela documents
        
        Args:
            file_info: Dicionário com informações do arquivo
            chapters: Lista de capítulos detectados
            document_id: ID do documento existente (para update)
            
        Returns:
            ID do documento inserido/atualizado ou None se erro
        """
        try:
            cursor = self.connection.cursor()
            
            if document_id:
                # Atualiza documento existente
                query = """
                UPDATE documents 
                SET file_name = %s, file_type = %s, file_size = %s, 
                    content_length = %s, file_hash = %s, modification_timestamp = %s,
                    chapters_count = %s, last_processed = CURRENT_TIMESTAMP, status = %s
                WHERE id = %s
                """
                
                values = (
                    file_info['name'],
                    file_info.get('extension', ''),
                    file_info.get('size', 0),
                    file_info.get('content_length', 0),
                    file_info['file_hash'],
                    file_info['modification_timestamp'],
                    len(chapters),
                    'processed',
                    document_id
                )
                
                cursor.execute(query, values)
                self.connection.commit()
                
                logger.info(f"Documento atualizado: {file_info['name']} ({len(chapters)} capítulos)")
                return document_id
                
            else:
                # Insere novo documento
                query = """
                INSERT INTO documents 
                (file_path, file_name, file_type, file_size, content_length, 
                 file_hash, modification_timestamp, chapters_count, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """
                
                values = (
                    file_info['path'],
                    file_info['name'],
                    file_info.get('extension', ''),
                    file_info.get('size', 0),
                    file_info.get('content_length', 0),
                    file_info['file_hash'],
                    file_info['modification_timestamp'],
                    len(chapters),
                    'processed'
                )
                
                cursor.execute(query, values)
                self.connection.commit()
                
                logger.info(f"Novo documento inserido: {file_info['name']} ({len(chapters)} capítulos)")
                return cursor.lastrowid
                
        except Error as e:
            logger.error(f"Erro ao inserir/atualizar documento: {e}")
            self.connection.rollback()
            return None
        finally:
            if cursor:
                cursor.close()
    
    def insert_chapters(self, document_id: int, chapters: List[Dict]) -> bool:
        """
        Insere informações dos capítulos na tabela document_chapters
        
        Args:
            document_id: ID do documento
            chapters: Lista de capítulos
            
        Returns:
            True se sucesso, False se erro
        """
        try:
            cursor = self.connection.cursor()
            
            # Remove capítulos existentes do documento
            cursor.execute("DELETE FROM document_chapters WHERE document_id = %s", (document_id,))
            
            if not chapters:
                self.connection.commit()
                return True
            
            # Insere novos capítulos
            query = """
            INSERT INTO document_chapters 
            (document_id, chapter_index, title, chapter_type, level, 
             start_position, end_position, content_preview)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """
            
            chapter_values = []
            for i, chapter in enumerate(chapters):
                chapter_values.append((
                    document_id,
                    i,
                    chapter['title'],
                    chapter['type'],
                    chapter['level'],
                    chapter['start_position'],
                    chapter.get('end_position', chapter['start_position'] + 1000),
                    chapter.get('content_preview', '')[:500]  # Limita preview
                ))
            
            cursor.executemany(query, chapter_values)
            self.connection.commit()
            
            logger.info(f"Inseridos {len(chapters)} capítulos para documento {document_id}")
            return True
            
        except Error as e:
            logger.error(f"Erro ao inserir capítulos: {e}")
            self.connection.rollback()
            return False
        finally:
            if cursor:
                cursor.close()
                
    def insert_chunks(self, document_id: int, chunks: List[Dict]) -> bool:
        """
        Insere chunks organizados por capítulos na tabela document_chunks
        
        Args:
            document_id: ID do documento
            chunks: Lista de chunks organizados por capítulos
            
        Returns:
            True se sucesso, False se erro
        """
        try:
            cursor = self.connection.cursor()
            
            # Remove chunks existentes do documento
            cursor.execute("DELETE FROM document_chunks WHERE document_id = %s", (document_id,))
            
            if not chunks:
                self.connection.commit()
                logger.warning(f"Nenhum chunk para inserir no documento {document_id}")
                return True
            
            # Insere novos chunks em lotes
            query = """
            INSERT INTO document_chunks 
            (document_id, chunk_index, content, chunk_size, start_position, end_position, 
             overlap_size, chapter_title, chapter_type, chapter_level, 
             absolute_start_position, absolute_end_position, is_chapter_complete)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """
            
            chunk_values = []
            for chunk in chunks:
                chunk_values.append((
                    document_id,
                    chunk['chunk_index'],
                    chunk['content'],
                    chunk['chunk_size'],
                    chunk['start_position'],
                    chunk['end_position'],
                    chunk['overlap_size'],
                    chunk['chapter_title'],
                    chunk['chapter_type'],
                    chunk['chapter_level'],
                    chunk['absolute_start_position'],
                    chunk['absolute_end_position'],
                    chunk['is_chapter_complete']
                ))
            
            # Insere em lotes de 100
            batch_size = 100
            for i in range(0, len(chunk_values), batch_size):
                batch = chunk_values[i:i + batch_size]
                cursor.executemany(query, batch)
            
            self.connection.commit()
            
            logger.info(f"Inseridos {len(chunks)} chunks organizados por capítulos para documento {document_id}")
            return True
            
        except Error as e:
            logger.error(f"Erro ao inserir chunks: {e}")
            self.connection.rollback()
            return False
        finally:
            if cursor:
                cursor.close()
                
    def log_processing(self, document_id: Optional[int], operation: str, 
                      status: str, message: str):
        """
        Registra log de processamento
        
        Args:
            document_id: ID do documento (pode ser None)
            operation: Operação realizada
            status: Status da operação (success, error, warning)
            message: Mensagem do log
        """
        try:
            cursor = self.connection.cursor()
            
            query = """
            INSERT INTO processing_logs (document_id, operation, status, message)
            VALUES (%s, %s, %s, %s)
            """
            
            cursor.execute(query, (document_id, operation, status, message))
            self.connection.commit()
            
        except Error as e:
            logger.error(f"Erro ao registrar log: {e}")
        finally:
            if cursor:
                cursor.close()
    
    def get_processing_stats(self) -> Dict:
        """
        Obtém estatísticas detalhadas do processamento do banco de dados
        
        Returns:
            Dicionário com estatísticas
        """
        try:
            cursor = self.connection.cursor()
            
            stats = {}
            
            # Estatísticas dos documentos
            cursor.execute("SELECT COUNT(*) FROM documents")
            stats['total_documents'] = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM documents WHERE status = 'processed'")
            stats['processed_documents'] = cursor.fetchone()[0]
            
            cursor.execute("SELECT SUM(chapters_count) FROM documents WHERE status = 'processed'")
            result = cursor.fetchone()
            stats['total_chapters'] = result[0] if result[0] else 0
            
            # Estatísticas dos chunks
            cursor.execute("SELECT COUNT(*) FROM document_chunks")
            stats['total_chunks'] = cursor.fetchone()[0]
            
            cursor.execute("SELECT AVG(chunk_size), MIN(chunk_size), MAX(chunk_size) FROM document_chunks")
            result = cursor.fetchone()
            stats['avg_chunk_size'] = round(result[0], 2) if result[0] else 0
            stats['min_chunk_size'] = result[1] if result[1] else 0
            stats['max_chunk_size'] = result[2] if result[2] else 0
            
            # Estatísticas por tipo de capítulo
            cursor.execute("""
                SELECT chapter_type, COUNT(*) 
                FROM document_chunks 
                GROUP BY chapter_type
                ORDER BY COUNT(*) DESC
            """)
            stats['chunks_by_chapter_type'] = dict(cursor.fetchall())
            
            # Tipos de arquivo
            cursor.execute("""
                SELECT file_type, COUNT(*) 
                FROM documents 
                GROUP BY file_type
                ORDER BY COUNT(*) DESC
            """)
            stats['file_types'] = dict(cursor.fetchall())
            
            return stats
            
        except Error as e:
            logger.error(f"Erro ao obter estatísticas: {e}")
            return {}
        finally:
            if cursor:
                cursor.close()

class LegalRAGProcessor:
    """Classe principal para processar documentos jurídicos com detecção de capítulos"""
    
    def __init__(self, db_config: Dict, documents_folder: str, gemini_api_key: str,
                 chunk_size: int = 1000, overlap: int = 200):
        """
        Inicializa o processador RAG para documentos jurídicos
        
        Args:
            db_config: Configurações do banco de dados
            documents_folder: Pasta com os documentos
            gemini_api_key: Chave da API do Google Gemini
            chunk_size: Tamanho dos chunks
            overlap: Sobreposição entre chunks
        """
        self.db_manager = LegalDatabaseManager(**db_config)
        self.file_processor = SmartLocalFileProcessor(documents_folder)
        self.chapter_detector = LegalDocumentChapterDetector(gemini_api_key)
        self.chunker = LegalDocumentChunker(chunk_size, overlap)
        self.text_extractor = EnhancedTextExtractor()
        self.documents_folder = documents_folder
        
    def process_documents(self, file_types: List[str] = None, 
                         recursive: bool = True,
                         max_depth: Optional[int] = None,
                         show_progress: bool = True) -> Dict:
        """
        Processa documentos jurídicos com detecção de capítulos
        
        Args:
            file_types: Tipos de arquivo para processar
            recursive: Se deve buscar em subpastas
            max_depth: Profundidade máxima de busca
            show_progress: Se deve mostrar progresso
            
        Returns:
            Dicionário com estatísticas do processamento
        """
        start_time = time.time()
        
        stats = {
            'total_files': 0,
            'new_files': 0,
            'updated_files': 0,
            'unchanged_files': 0,
            'failed_files': 0,
            'total_chapters': 0,
            'total_chunks': 0,
            'processing_time': 0,
            'folders_scanned': 0,
            'files_by_type': {},
            'chapters_by_type': {},
            'errors': []
        }
        
        try:
            # Conecta ao banco de dados
            self.db_manager.connect()
            
            logger.info("Iniciando busca recursiva de documentos jurídicos...")
            
            # Lista arquivos da pasta local
            files = self.file_processor.list_files(
                file_types=file_types, 
                recursive=recursive,
                max_depth=max_depth,
                show_progress=show_progress
            )
            
            stats['total_files'] = len(files)
            stats['folders_scanned'] = self.file_processor.search_stats['folders_scanned']
            stats['files_by_type'] = self.file_processor.search_stats['files_by_type'].copy()
            
            if not files:
                logger.warning("Nenhum arquivo encontrado para processar!")
                return stats
            
            logger.info(f"Iniciando processamento de {len(files)} documentos jurídicos...")
            
            # Processa cada arquivo
            for i, file_info in enumerate(files, 1):
                try:
                    if show_progress and i % 2 == 0:  # Mostra progresso mais frequentemente
                        logger.info(f"Processando documento {i}/{len(files)}: {file_info['name']}")
                    
                    # Verifica status do arquivo no banco
                    file_status = self.db_manager.check_file_status(file_info)
                    
                    if file_status['exists'] and not file_status['needs_update']:
                        # Arquivo existe e não foi modificado
                        logger.debug(f"Documento inalterado, pulando: {file_info['name']}")
                        stats['unchanged_files'] += 1
                        continue
                    
                    # Processa documento (novo ou modificado)
                    result = self._process_single_document(file_info, file_status)
                    
                    if result:
                        chapters_count, chunks_count = result
                        
                        if file_status['exists']:
                            stats['updated_files'] += 1
                            logger.info(f"Documento atualizado: {file_info['name']} "
                                      f"({chapters_count} capítulos, {chunks_count} chunks)")
                        else:
                            stats['new_files'] += 1
                            logger.info(f"Novo documento processado: {file_info['name']} "
                                      f"({chapters_count} capítulos, {chunks_count} chunks)")
                        
                        stats['total_chapters'] += chapters_count
                        stats['total_chunks'] += chunks_count
                    else:
                        stats['failed_files'] += 1
                        
                except Exception as e:
                    error_msg = f"Erro ao processar documento {file_info['name']}: {e}"
                    logger.error(error_msg)
                    stats['failed_files'] += 1
                    stats['errors'].append(error_msg)
                    
        except Exception as e:
            error_msg = f"Erro geral no processamento: {e}"
            logger.error(error_msg)
            stats['errors'].append(error_msg)
            
        finally:
            self.db_manager.disconnect()
            stats['processing_time'] = round(time.time() - start_time, 2)
            
        return stats
    
    def _process_single_document(self, file_info: Dict, file_status: Dict) -> Optional[Tuple[int, int]]:
        """
        Processa um único documento jurídico
        
        Args:
            file_info: Informações do arquivo
            file_status: Status do arquivo no banco
            
        Returns:
            Tupla com (número de capítulos, número de chunks) ou None se erro
        """
        try:
            logger.info(f"Processando documento jurídico: {file_info['relative_path']}")
            
            # Extrai texto do documento
            file_extension = file_info['extension'].lower()
            file_path = file_info['path']
            
            text = self.text_extractor.extract_text(file_path, file_extension)
            
            if not text or len(text.strip()) == 0:
                self.db_manager.log_processing(
                    file_status.get('document_id'), 'text_extraction', 'error',
                    f"Não foi possível extrair texto de {file_info['name']}"
                )
                logger.warning(f"Nenhum texto extraído de {file_info['name']}")
                return None
            
            # Adiciona informações do conteúdo
            file_info['content_length'] = len(text)
            
            # Detecta capítulos usando IA
            logger.info(f"Detectando capítulos em {file_info['name']}...")
            chapters = self.chapter_detector.detect_chapters(text, file_info['name'])
            
            if not chapters:
                logger.warning(f"Nenhum capítulo detectado em {file_info['name']}")
            
            # Insere ou atualiza documento no banco
            document_id = self.db_manager.insert_or_update_document(
                file_info, 
                chapters,
                file_status.get('document_id')
            )
            
            if not document_id:
                return None
            
            # Insere informações dos capítulos
            if chapters:
                success = self.db_manager.insert_chapters(document_id, chapters)
                if not success:
                    logger.error(f"Erro ao inserir capítulos do documento {file_info['name']}")
            
            # Cria chunks organizados por capítulos
            logger.info(f"Criando chunks por capítulos para {file_info['name']}...")
            chunks = self.chunker.create_chapter_chunks(text, chapters)
            
            if not chunks:
                self.db_manager.log_processing(
                    document_id, 'chunking', 'warning',
                    f"Nenhum chunk criado para {file_info['name']}"
                )
                logger.warning(f"Nenhum chunk criado para {file_info['name']}")
                return (len(chapters), 0)
            
            # Insere chunks no banco
            success = self.db_manager.insert_chunks(document_id, chunks)
            
            if success:
                operation = 'update' if file_status['exists'] else 'insert'
                self.db_manager.log_processing(
                    document_id, f'legal_processing_{operation}', 'success',
                    f"Processado documento jurídico: {len(chapters)} capítulos, {len(chunks)} chunks"
                )
                return (len(chapters), len(chunks))
            else:
                return None
                
        except Exception as e:
            logger.error(f"Erro no processamento do documento {file_info['name']}: {e}")
            if 'document_id' in locals():
                self.db_manager.log_processing(
                    document_id, 'legal_processing', 'error', str(e)
                )
            return None
    
    def get_processing_stats(self) -> Dict:
        """
        Obtém estatísticas detalhadas do processamento
        
        Returns:
            Dicionário com estatísticas
        """
        try:
            self.db_manager.connect()
            return self.db_manager.get_processing_stats()
        finally:
            self.db_manager.disconnect()

# Schema SQL atualizado para documentos jurídicos
def create_legal_database_schema():
    """
    Retorna o script SQL para documentos jurídicos com capítulos
    """
    return """
    -- Tabela para armazenar informações dos documentos jurídicos
    CREATE TABLE IF NOT EXISTS documents (
        id INT AUTO_INCREMENT PRIMARY KEY,
        file_path VARCHAR(1000) UNIQUE NOT NULL,
        file_name VARCHAR(500) NOT NULL,
        file_type VARCHAR(50) NOT NULL,
        file_size BIGINT,
        content_length INT,
        file_hash VARCHAR(32) NOT NULL,
        modification_timestamp DOUBLE NOT NULL,
        chapters_count INT DEFAULT 0,
        upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        last_processed TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
        status ENUM('pending', 'processed', 'error') DEFAULT 'pending',
        
        INDEX idx_file_path (file_path),
        INDEX idx_file_hash (file_hash),
        INDEX idx_modification_timestamp (modification_timestamp),
        INDEX idx_status (status),
        INDEX idx_chapters_count (chapters_count)
    );

    -- Tabela para armazenar informações dos capítulos detectados
    CREATE TABLE IF NOT EXISTS document_chapters (
        id INT AUTO_INCREMENT PRIMARY KEY,
        document_id INT NOT NULL,
        chapter_index INT NOT NULL,
        title VARCHAR(500) NOT NULL,
        chapter_type ENUM('CAPITULO', 'SECAO', 'SUBSECAO', 'ARTIGO', 'DISPOSITIVO', 'FUNDAMENTACAO', 'TITULO', 'DOCUMENTO') NOT NULL,
        level INT NOT NULL,
        start_position INT NOT NULL,
        end_position INT,
        content_preview TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        
        FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE CASCADE,
        INDEX idx_document_chapters (document_id, chapter_index),
        INDEX idx_chapter_type (chapter_type),
        INDEX idx_chapter_level (level),
        FULLTEXT INDEX idx_chapter_title (title)
    );

    -- Tabela para armazenar os chunks organizados por capítulos
    CREATE TABLE IF NOT EXISTS document_chunks (
        id INT AUTO_INCREMENT PRIMARY KEY,
        document_id INT NOT NULL,
        chunk_index INT NOT NULL,
        content TEXT NOT NULL,
        chunk_size INT NOT NULL,
        start_position INT NOT NULL,
        end_position INT NOT NULL,
        overlap_size INT DEFAULT 0,
        chapter_title VARCHAR(500),
        chapter_type ENUM('CAPITULO', 'SECAO', 'SUBSECAO', 'ARTIGO', 'DISPOSITIVO', 'FUNDAMENTACAO', 'TITULO', 'DOCUMENTO'),
        chapter_level INT,
        absolute_start_position INT,
        absolute_end_position INT,
        is_chapter_complete BOOLEAN DEFAULT FALSE,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        
        FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE CASCADE,
        INDEX idx_document_id (document_id),
        INDEX idx_chunk_index (document_id, chunk_index),
        INDEX idx_chapter_info (chapter_type, chapter_level),
        FULLTEXT INDEX idx_content (content),
        FULLTEXT INDEX idx_chapter_title (chapter_title)
    );

    -- Tabela para armazenar metadados adicionais dos chunks
    CREATE TABLE IF NOT EXISTS chunk_metadata (
        id INT AUTO_INCREMENT PRIMARY KEY,
        chunk_id INT NOT NULL,
        metadata_key VARCHAR(100) NOT NULL,
        metadata_value TEXT,
        
        FOREIGN KEY (chunk_id) REFERENCES document_chunks(id) ON DELETE CASCADE,
        INDEX idx_chunk_metadata (chunk_id, metadata_key)
    );

    -- Tabela para logs de processamento
    CREATE TABLE IF NOT EXISTS processing_logs (
        id INT AUTO_INCREMENT PRIMARY KEY,
        document_id INT,
        operation VARCHAR(100) NOT NULL,
        status ENUM('success', 'error', 'warning') NOT NULL,
        message TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        
        FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE SET NULL,
        INDEX idx_document_logs (document_id),
        INDEX idx_operation (operation),
        INDEX idx_status (status)
    );
    """

# Exemplo de uso principal
if __name__ == "__main__":
    # Configurações do banco de dados
    db_config = {
        'host': 'localhost',
        'database': 'legal_rag_database',
        'user': 'seu_usuario',
        'password': 'sua_senha',
        'port': 3306
    }
    
    # Configurações
    documents_folder = "C:/Caminho/Para/Documentos/Juridicos"
    gemini_api_key = "SUA_CHAVE_API_GEMINI"  # Obtenha em https://makersuite.google.com/app/apikey

    try:
        # Inicializa o processador de documentos jurídicos
        processor = LegalRAGProcessor(
            db_config=db_config,
            documents_folder=documents_folder,
            gemini_api_key=gemini_api_key,
            chunk_size=1000,  # Tamanho do chunk em caracteres
            overlap=200       # Sobreposição em caracteres
        )
        
        print("\n" + "="*80)
        print("SISTEMA DE PROCESSAMENTO DE DOCUMENTOS JURÍDICOS")
        print("="*80)
        print("Funcionalidades:")
        print("- Detecção automática de capítulos usando IA (Gemini)")
        print("- Chunking organizado por capítulos")
        print("- Suporte a despachos, decisões e sentenças")
        print("- Detecção inteligente de mudanças")
        print("- Estrutura hierárquica de documentos")
        
        # Processa documentos jurídicos
        stats = processor.process_documents(
            file_types=['pdf', 'docx', 'txt'],  # Tipos de arquivo
            recursive=True,                      # Busca recursiva
            max_depth=None,                      # Sem limite de profundidade
            show_progress=True                   # Mostrar progresso
        )
        
        # Exibe estatísticas detalhadas
        print("\n" + "="*80)
        print("ESTATÍSTICAS DO PROCESSAMENTO DE DOCUMENTOS JURÍDICOS")
        print("="*80)
        print(f"Tempo total de processamento: {stats['processing_time']} segundos")
        print(f"Pastas escaneadas: {stats['folders_scanned']}")
        print(f"Total de documentos encontrados: {stats['total_files']}")
        print(f"Documentos novos processados: {stats['new_files']}")
        print(f"Documentos atualizados: {stats['updated_files']}")
        print(f"Documentos inalterados (pulados): {stats['unchanged_files']}")
        print(f"Documentos com erro: {stats['failed_files']}")
        print(f"Total de capítulos detectados: {stats['total_chapters']}")
        print(f"Total de chunks criados: {stats['total_chunks']}")
        
        if stats['files_by_type']:
            print(f"\nDocumentos por tipo:")
            for file_type, count in stats['files_by_type'].items():
                if count > 0:
                    print(f"  - .{file_type}: {count} documentos")
        
        if stats['errors']:
            print(f"\nErros encontrados ({len(stats['errors'])}):")
            for error in stats['errors'][:3]:  # Mostra apenas os primeiros 3
                print(f"  - {error}")
            if len(stats['errors']) > 3:
                print(f"  ... e mais {len(stats['errors']) - 3} erros")
        
        # Obtém e exibe estatísticas do banco de dados
        print("\n" + "="*80)
        print("ESTATÍSTICAS DO BANCO DE DADOS JURÍDICOS")
        print("="*80)
        
        db_stats = processor.get_processing_stats()
        if db_stats:
            print(f"Total de documentos no banco: {db_stats['total_documents']}")
            print(f"Documentos processados: {db_stats['processed_documents']}")
            print(f"Total de capítulos no banco: {db_stats['total_chapters']}")
            print(f"Total de chunks no banco: {db_stats['total_chunks']}")
            print(f"Tamanho médio dos chunks: {db_stats['avg_chunk_size']} caracteres")
            print(f"Menor chunk: {db_stats['min_chunk_size']} caracteres")
            print(f"Maior chunk: {db_stats['max_chunk_size']} caracteres")
            
            if db_stats['chunks_by_chapter_type']:
                print(f"\nChunks por tipo de capítulo:")
                for chapter_type, count in db_stats['chunks_by_chapter_type'].items():
                    print(f"  - {chapter_type}: {count} chunks")
            
            if db_stats['file_types']:
                print(f"\nTipos de arquivo no banco:")
                for file_type, count in db_stats['file_types'].items():
                    print(f"  - {file_type}: {count} documentos")
        
        print("="*80)
        print("PROCESSAMENTO DE DOCUMENTOS JURÍDICOS CONCLUÍDO!")
        print("="*80)
        print("\nRecursos implementados:")
        print("✓ Detecção automática de capítulos com IA")
        print("✓ Chunking organizado por estrutura jurídica")
        print("✓ Suporte a despachos, decisões e sentenças")
        print("✓ Detecção inteligente de mudanças")
        print("✓ Busca recursiva em subpastas")
        print("✓ Logs detalhados de processamento")
        print("\nPróximas execuções processarão apenas documentos novos ou modificados!")
        
    except Exception as e:
        logger.error(f"Erro fatal no processamento: {e}")
        print(f"\nERRO FATAL: {e}")
        print("Verifique as configurações e tente novamente.")
        print("\nVerifique se:")
        print("1. A chave da API do Gemini está correta")
        print("2. O banco de dados está acessível")
        print("3. A pasta de documentos existe")
        print("4. As dependências estão instaladas:")
        print("   pip install google-generativeai mysql-connector-python PyPDF2 python-docx")

    # Exibe o schema SQL para documentos jurídicos
    print("\n" + "="*80)
    print("SCHEMA SQL PARA DOCUMENTOS JURÍDICOS")
    print("="*80)
    print("Execute este SQL no seu banco de dados:")
    print(create_legal_database_schema())
